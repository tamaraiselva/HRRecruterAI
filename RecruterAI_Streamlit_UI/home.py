import streamlit as st
from modal import Modal
import streamlit_antd_components as sac
import requests
from pymongo import MongoClient
# from streamlit_extras.switch_page_button import switch_page
import pandas as pd
import re
import time as t
import uuid
import ast
import numpy as np 
import streamlit_shadcn_ui as ui
from datetime import datetime, date, time
from emailconfig_User import EmailConfig_User
import pymongo
import time
import io
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import os
import time
import pythoncom
import pymongo
import logging
from dotenv import load_dotenv

load_dotenv()

# Fetch the API key from environment variables
authorization_key = os.getenv("BLANDAI_API_KEY")



# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def main():
    # MongoDB configuration
    if "message" in st.session_state and st.session_state.message:
        st.success(st.session_state.message)
        st.session_state.message = ""

    MONGO_URI = "mongodb://localhost:27017/"
    client = MongoClient(MONGO_URI)
    db = client["recruiter_ai"]
    collection = db["job_descriptions"]
    candidate_collection = db["candidates"]
    manage_job = db["manage_job"]
    calls_collection = db["call_conversations"]

    st.title("HR Recruiter AI")

    # if 'logged_in' not in st.session_state or not st.session_state.logged_in:
    #     switch_page("login")
    
    # Define modal
    modal = Modal("Job Description", key="Job_Description", max_width=1000, padding=20,top_position="10px")

    # Initialize modal for managing users
    manage_users_modal = Modal("Manage Users", key="Manage_Users", max_width=1000, padding=20,top_position="10px")
    edit_user_modal = Modal("Edit User", key="Edit_User", max_width=1000, padding=20,top_position="10px")
    add_user_modal = Modal("Add User", key="Add_User", max_width=1000, padding=20,top_position="10px")
    settings_model = Modal("Settings", key="Settings", max_width=1000, padding=20,top_position="10px")

    manage_job_modal = Modal("Manage Job", key="manage_job_modal", max_width=1000, padding=20,top_position="10px")
    add_job_modal = Modal("Add Panel for the Round", key="add_job_modal", max_width=1000, padding=20,top_position="10px")
    edit_job_modal = Modal("Edit Member", key="edit_job_modal", max_width=1000, padding=20,top_position="10px")

    # Initialize modal for managing members
    manage_members_modal = Modal("Manage Members", key="Manage_Members", max_width=1000, padding=20,top_position="10px")
    edit_members_modal = Modal("Edit Members", key="Edit_Members", max_width=1000, padding=20,top_position="10px")
    add_members_modal = Modal("Add Members", key="Add_Members", max_width=1000, padding=20,top_position="10px")

    # Initialize session state variables
    if 'current_job_description' not in st.session_state:
        st.session_state['current_job_description'] = ""
    if 'selected_job_id' not in st.session_state:
        st.session_state['selected_job_id'] = None
    if "selected_candidate" not in st.session_state:
        st.session_state["selected_candidate"] = None
    if 'modal_open' not in st.session_state:
        st.session_state['modal_open'] = False
    if 'modal_content' not in st.session_state:
        st.session_state['modal_content'] = ""
    if 'job_submitted' not in st.session_state:
        st.session_state['job_submitted'] = False
    if 'job_updated' not in st.session_state:
        st.session_state['job_updated'] = False
    if 'JD_success_flag' not in st.session_state:
        st.session_state['JD_success_flag'] = False
    if 'setting_mobile_number_flag' not in st.session_state:
        st.session_state['setting_mobile_number_flag'] = False
    if 'setting_flag' not in st.session_state:
        st.session_state['setting_flag'] = False
    if 'whatsapp_success_flag' not in st.session_state:
        st.session_state['whatsapp_success_flag'] = False
    if 'JD_retrieve_error_flag' not in st.session_state:
        st.session_state['JD_retrieve_error_flag'] = False
    if 'JD_bad_request_error_flag' not in st.session_state:
        st.session_state['JD_bad_request_error_flag'] = False
    if 'JD_not_found_flag' not in st.session_state:
        st.session_state['JD_not_found_flag'] = False
    if 'JD_warning_flag' not in st.session_state:
        st.session_state['JD_warning_flag'] = False
    if 'update_success_flag' not in st.session_state:
        st.session_state['update_success_flag'] = False
    if 'update_error_flag' not in st.session_state:
        st.session_state['update_error_flag'] = False
    if 'update_fetch_error_flag' not in st.session_state:
        st.session_state['update_fetch_error_flag'] = False
    if 'update_warning_flag' not in st.session_state:
        st.session_state['update_warning_flag'] = False
    if 'view_candidates' not in st.session_state:
        st.session_state['view_candidates'] = False
    if "selected_job_id" not in st.session_state:
        st.session_state["selected_job_id"] = None
    if 'creating_new_job' not in st.session_state:
        st.session_state['creating_new_job'] = False

    if 'checked_candidates' not in st.session_state:
        st.session_state['checked_candidates'] = set()  # Initialize as a set for fast add/remove operations
    if 'status_update' not in st.session_state:
        st.session_state['status_update'] = {}
    if 'df_candidates' not in st.session_state:
        st.session_state['df_candidates'] = pd.DataFrame()  # Initialize with empty DataFrame

    if 'call_queued_flag' not in st.session_state:
        st.session_state['call_queued_flag'] = False
    if 'interested' not in st.session_state:
        st.session_state['interested'] = False

    if 'candidate_status' not in st.session_state:
        st.session_state['candidate_status'] = ''
    

    if 'user_id' not in st.session_state:
        st.session_state['user_id'] = None
    if 'user_role' not in st.session_state:
        st.session_state['user_role'] = None
    if 'show_view_candidates' not in st.session_state:
        st.session_state['show_view_candidates'] = False
    if 'save_clicked' not in st.session_state:
        st.session_state['save_clicked'] = False
    if 'first_save_clicked' not in st.session_state:
        st.session_state['first_save_clicked'] = False 
    if 'selected_candidate_id' not in st.session_state:
        st.session_state['selected_candidate_id'] = None  

    if 'edit_mode' not in st.session_state:
        st.session_state['edit_mode'] = False
    if 'update_mode' not in st.session_state:
        st.session_state['update_mode'] = False

    if 'modal_candidate_id' not in st.session_state:
        st.session_state['modal_candidate_id'] = None
    if 'modal_candidate_name' not in st.session_state:
        st.session_state['modal_candidate_name'] = None

    if 'edit_user_mode' not in st.session_state:
        st.session_state['edit_user_mode'] = False
    if 'edit_manage_job_mode' not in st.session_state:
        st.session_state['edit_manage_job_mode'] = False
    if 'edit_member_mode' not in st.session_state:
        st.session_state['edit_member_mode'] = False
    if 'button_colors' not in st.session_state:
        st.session_state['button_colors'] = {}

    # Alert boxes
    # Display success message at the top if the flag is set
    if st.session_state['setting_flag']:
        st.error("Please go to settings and update the WhatsApp API Token and WhatsApp Cloud Number ID.")
        st.session_state['setting_flag'] = False  # Reset the flag

    # Display success message at the top if the flag is set
    if st.session_state['setting_mobile_number_flag']:
        st.error("Mobile number must be exactly 10 digits.")
        st.session_state['setting_mobile_number_flag'] = False  # Reset the flag
        
    # Display success message at the top if the flag is set
    if st.session_state['whatsapp_success_flag']:
        st.success(f"Message sent successfully for the selected Candidate....")
        st.session_state['whatsapp_success_flag'] = False  # Reset the flag

    # Display success message at the top if the flag is set
    if st.session_state['JD_success_flag']:
        st.success("Job description created successfully!")
        st.session_state['JD_success_flag'] = False  # Reset the flag

    if st.session_state['JD_retrieve_error_flag']:
        st.error("Failed to retrieve job ID or job description from response.")
        st.session_state['JD_retrieve_error_flag'] = False

    if st.session_state['JD_bad_request_error_flag']:
        st.error("Bad Error Key from Hugging Face......")
        st.session_state['JD_bad_request_error_flag'] = False

    if st.session_state['JD_not_found_flag']:
        st.error("Job Description not found.")
        st.session_state['JD_not_found_flag'] = False

    if st.session_state['JD_warning_flag']:
        st.warning("Please enter a job description before submitting.")
        st.session_state['JD_warning_flag'] = False

    if st.session_state['call_queued_flag']:
        st.success("Call queued successfully for the selected Candidate(s).")
        st.session_state['call_queued_flag'] = False
    
    # Display success message if a candidate is marked as "Interested"
    if st.session_state['interested']:
        st.success("The candidate has shown interest based on the latest call.")
        st.session_state['interested'] = False  # Reset the flag if necessary

    if 'delete_confirmation' not in st.session_state:
        st.session_state['delete_confirmation'] = False

    

    # Fetch job descriptions based on user role
    def fetch_job_descriptions():
        if st.session_state.user_role == 'admin':
            return list(collection.find({}, {"_id": 1, "prompt": 1, "job_description": 1}).sort([('_id', -1)]))
        else:
            user_id = st.session_state.get('user_id')
            return list(collection.find({"user_id": user_id}, {"_id": 1, "prompt": 1, "job_description": 1}).sort([('_id', -1)]))

    # Define validation functions
    def validate_name(name):
        return re.match(r"^[A-Z][a-zA-Z0-9\s]*$", name)
    
    def validate_email(email):
        return re.match(r"[^@]+@[^@]+\.[^@]+", email)
    
    def validate_password(password):
        return re.match(r"^(?=.*[A-Z])(?=.*[a-z])(?=.*\d).{8,}$", password)
    
    def validate_mobile_number(mobile_number):
        return re.match(r"^\d{10}$", mobile_number)
    
    # Function to show the delete confirmation modal
    # @st.experimental_dialog("Delete Confirmation")
    @st.experimental_dialog("Delete Confirmation")
    def Delete_Panel(panel):
        st.write(f"Are you sure you want to delete {panel['name']}?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                delete_job_panel(st.session_state['selected_job_id'], panel['id'])
                st.session_state.delete_success = True
                st.rerun()
        with col2:
            if st.button("No"):
                st.rerun()

    # Function to show the delete confirmation modal
    # @st.experimental_dialog("Delete Confirmation")
    @st.experimental_dialog("Delete Confirmation")
    def Delete_Users(user):
        st.write(f"Are you sure you want to delete {user['name']}?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                delete_user(user['id'])
                st.session_state.delete_success = True
                st.rerun()
        with col2:
            if st.button("No"):
                st.rerun()

    # Function to show the delete confirmation modal
    # @st.experimental_dialog("Delete Confirmation")
    @st.experimental_dialog("Delete Confirmation")
    def Delete_Members(member):
        st.write(f"Are you sure you want to delete {member['name']}?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                delete_member(member['id'])
                st.session_state.delete_success = True
                st.rerun()
        with col2:
            if st.button("No"):
                st.rerun()

    # Handle New Job Description
    def new_job_description():
        st.session_state['current_job_description'] = ""
        st.session_state['selected_job_id'] = None
        st.session_state['job_submitted'] = False
        st.session_state['job_updated'] = False
        st.session_state['creating_new_job'] = True
        st.session_state['view_candidates'] = False
        st.session_state['edit_mode'] = False
        st.session_state['job_updated'] = False
        st.session_state['checked_candidates'] = set()

    # Function to generate PDF from job description
    def generate_pdf(job_description):
         # Initialize COM
        pythoncom.CoInitialize()
        # Create a new Word document
        doc = Document()

        # Add a title to the document
        title = doc.add_heading("Job Description", level=1)
        title_format = title.runs[0].font
        title_format.size = Pt(24)  # Set title font size

        # Split the job description into lines and add to the document
        lines = job_description.split("\n")
        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                # Add each line to the document as a paragraph
                paragraph = doc.add_paragraph(stripped_line)
                paragraph_format = paragraph.runs[0].font
                paragraph_format.size = Pt(12)  # Set paragraph font size

        # Save the document to a temporary .docx file in memory
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        # Write the .docx content to a temporary file
        with open("temp_job_description.docx", "wb") as temp_file:
            temp_file.write(docx_buffer.getbuffer())

        # Convert the .docx file to a PDF
        convert("temp_job_description.docx", "temp_job_description.pdf")

        # Read the PDF content into a BytesIO buffer
        pdf_buffer = io.BytesIO()
        with open("temp_job_description.pdf", "rb") as pdf_file:
            pdf_buffer.write(pdf_file.read())

        # Clean up temporary files
        os.remove("temp_job_description.docx")
        os.remove("temp_job_description.pdf")

        # Return the PDF buffer
        pdf_buffer.seek(0)
        return pdf_buffer
        
    # Function to show the PDF download dialog
    @st.experimental_dialog("Download PDF")
    def download_pdf_dialog(job_description):
        st.write("Click below to download the job description as a PDF:")
        with st.spinner('Loading...'):
            pdf_buffer = generate_pdf(job_description)
            st.download_button(
                label="Download PDF",
                data=pdf_buffer,
                file_name=f"Job_ID_{st.session_state['selected_job_id']}.pdf",
                mime="application/pdf"
            )

    # Function to handle the download logic
    def handle_download():
        if st.session_state.get('selected_job_id') is not None:
            api_url = f"http://localhost:8081/api/v1/jd/{st.session_state['selected_job_id']}"
            response = requests.get(api_url)
            if response.status_code == 200:
                jd_response = response.json()
                job_description = jd_response.get('job_description', '')
                download_pdf_dialog(job_description)  # Open the dialog
            else:
                st.error("Failed to fetch job description for download.")

    # Function to show the confirmation dialog for editing the job description
    @st.experimental_dialog("Confirm Edit")
    def confirm_edit_dialog():
        st.write("Do you want to edit the job description?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                st.session_state['edit_mode'] = True
                st.session_state['save_clicked'] = True  # Disable Save button
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                st.rerun()
        with col2:
            if st.button("No"):
                st.session_state['edit_mode'] = False
                st.session_state['save_clicked'] = False  # Enable Save button
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                st.rerun()

    # Function to confirm editing for users
    @st.experimental_dialog("Confirm Edit for Users")
    def confirm_edit_user_dialog():
        st.write("Do you want to edit the user details?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                st.session_state['edit_user_mode'] = True
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                edit_user_modal.open()
                st.rerun()
        with col2:
            if st.button("No"):
                st.session_state['edit_user_mode'] = False
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                st.rerun()

    # Function to confirm editing for Manage job
    @st.experimental_dialog("Confirm Edit for Manage Job")
    def confirm_edit_manage_job_dialog():
        st.write("Do you want to edit the manage job details?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                st.session_state['edit_manage_job_mode'] = True
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                edit_job_modal.open()
                st.rerun()
        with col2:
            if st.button("No"):
                st.session_state['edit_manage_job_mode'] = False
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                st.rerun()

    # Function to confirm editing for users
    @st.experimental_dialog("Confirm Edit for Manage Job")
    def confirm_edit_member_dialog():
        st.write("Do you want to edit the Member details?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Yes"):
                st.session_state['edit_member_mode'] = True
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                edit_members_modal.open()
                st.rerun()
        with col2:
            if st.button("No"):
                st.session_state['edit_member_mode'] = False
                st.session_state['show_edit_dialog'] = False  # Close the dialog
                st.rerun()
                
    def submit_job_description(job_description):
        if job_description:
            api_url = "http://localhost:8081/api/v1/jd"
            payload = {"prompt": job_description}
            with st.spinner('Generating job description...'):
                try:
                    response = requests.post(api_url, json=payload)
                    if response.status_code == 201:
                        jd_response = response.json()
                        job_id = jd_response.get("id")
                        prompt_saved = job_description
                        job_description_created = jd_response.get("job_description")
                        user_id = st.session_state['user_id']  # Get the user_id from the session state

                        if job_id and job_description_created:
                            collection.insert_one({"_id": job_id, "prompt": prompt_saved, "job_description": job_description_created, "user_id": user_id})
                            st.session_state['selected_job_id'] = job_id
                            st.session_state['current_job_description'] = prompt_saved
                            st.session_state['job_submitted'] = True
                            st.session_state['JD_success_flag'] = True
                            st.session_state['save_clicked'] = False
                            st.session_state['first_save_clicked'] = False
                            st.rerun()
                        else:
                            st.session_state['JD_retrieve_error_flag'] = True
                    else:
                        if response.status_code == 500:
                            st.session_state['JD_bad_request_error_flag'] = True
                        elif response.status_code == 404:
                            st.session_state['JD_not_found_flag'] = True

                except requests.exceptions.RequestException as e:
                    st.error("Error")  # Display a generic error message
                    st.session_state['JD_retrieve_error_flag'] = True
        else:
            # st.warning("Please provide a job description.")
            st.session_state['JD_warning_flag'] = True


    # Update Job Description
    def update_job_description(job_description):
        if st.session_state['update_success_flag']:
            st.success("Job description updated successfully!")
            st.session_state['update_success_flag'] = False

        if st.session_state['update_error_flag']:
            st.error("Failed to update job description. Please try again.")
            st.session_state['update_error_flag'] = False

        if st.session_state['update_fetch_error_flag']:
            st.error("Failed to fetch job description. Please try again.")
            st.session_state['update_fetch_error_flag'] = False

        if st.session_state['update_warning_flag']:
            st.warning("No job description selected.")
            st.session_state['update_warning_flag'] = False
        with st.spinner('Updating job description...'):
            if st.session_state['selected_job_id'] is not None:
                api_url = f"http://localhost:8081/api/v1/jd/{st.session_state['selected_job_id']}"
                payload = {"job_description": job_description}
                update_response = requests.put(api_url, json=payload)

                if update_response.status_code == 200:
                    st.session_state['update_success_flag'] = True
                    st.session_state['job_updated'] = False
                    st.session_state['edit_mode'] = False
                    st.session_state['update_mode'] = False
                    st.session_state['first_save_clicked'] = False  # Reset this to allow re-enabling Save button
                    st.session_state['save_clicked'] = False
                    # Update the job description in MongoDB
                    collection.update_one(
                        {"_id": st.session_state['selected_job_id']},
                        {"$set": {"job_description": job_description}}
                    )
                    
                    st.rerun()
                else:
                    st.session_state['update_error_flag'] = True
            else:
                st.session_state['update_warning_flag'] = True

    
    candidate_model = Modal(
        "Candidate Status",
        key="Candidates",
        max_width=1250,
        padding=20,
        top_position="10px"
    )

    def display_candidates():
        with st.spinner('Loading Candidates...'):
            col1, col2, col3, col4, col5, col6 = st.columns([3, 2, 5, 6, 6, 6])
            with col1:
                st.markdown("**Select**")
            with col2:
                st.markdown("**Id**")
            with col3:
                st.markdown("**Name**")
            with col4:
                st.markdown("**Email**")
            with col5:
                st.markdown("**Mobile**")
            with col6:
                st.markdown("**Status**")

            selected_job_id = st.session_state.get('selected_job_id')
            job_description = collection.find_one({"_id": selected_job_id})


            if job_description and "resource_id" in job_description:
                resource_id = job_description["resource_id"]
                candidates = list(candidate_collection.find({"resource_id": resource_id, "selected_job_id": selected_job_id}) or [])

                candidate_list = []
                for candidate in candidates:
                    candidate_dict = {
                        "ID": candidate.get("ID"),
                        "Name": candidate.get("Name"),
                        "Email": candidate.get("Email"),
                        "Phone": candidate.get("Phone"),
                        "Status": candidate.get("Status")
                    }
                    candidate_list.append(candidate_dict)

                df_candidates = pd.DataFrame(candidate_list)
                st.session_state['df_candidates'] = df_candidates  # Store in session state

                if 'checked_candidates' not in st.session_state:
                    st.session_state['checked_candidates'] = set()
                if 'status_update' not in st.session_state:
                    st.session_state['status_update'] = {}
                if not df_candidates.empty:
                    for idx, row in df_candidates.iterrows():
                        col1, col2, col3, col4, col5, col6 = st.columns([3, 2, 5, 6, 6, 6])

                        with col1:
                            selected = st.checkbox("", key=f"candidate_{idx}", value=idx in st.session_state['checked_candidates'])

                            if selected:
                                st.session_state['checked_candidates'].add(idx)
                            else:
                                st.session_state['checked_candidates'].discard(idx)
                        with col2:
                            st.write(row['ID'])
                        with col3:
                            st.write(row['Name'])
                        with col4:
                            st.write(row['Email'])
                        with col5:
                            st.write(row['Phone'])
                        with col6:
                            if st.button("Status", key=f"change_status_{idx}"):
                                st.session_state['selected_candidate_id'] = row['ID']
                                candidate_model.open()

                    col7, col8 = st.columns(2)
                    with col7:
                        if st.button("Schedule Interview"):
                            if st.session_state['checked_candidates']:
                                selected_job_id = st.session_state['selected_job_id']
                                job_description = collection.find_one({"_id": selected_job_id})
                                if job_description and "resource_id" in job_description:
                                    resource_id = job_description["resource_id"]
                                    schedule_interviews()
                                    st.session_state['checked_candidates'] = set()
                                    st.session_state['status_update'] = {}
                                    st.rerun()
                                else:
                                    st.error("Failed to fetch job description.")
                            else:
                                st.warning("Please select a candidate for scheduling the interview.")
                    with col8:
                        ignored_candidates_button = st.button("Ignore Candidates")
                        if ignored_candidates_button:
                            selected_job_id = st.session_state['selected_job_id']
                            job_description = collection.find_one({"_id": selected_job_id})

                            if job_description and "resource_id" in job_description:
                                resource_id = job_description["resource_id"]

                                # Get the IDs of the candidates that have been selected (checked) for ignoring
                                selected_candidates = df_candidates.iloc[list(st.session_state['checked_candidates'])]
                                ignored_ids = selected_candidates["ID"].tolist()

                                # Delete the selected candidates from the database
                                candidate_collection.delete_many({"ID": {"$in": ignored_ids}, "resource_id": resource_id, "selected_job_id": selected_job_id})

                                # Update the DataFrame by removing ignored candidates
                                df_candidates = df_candidates.drop(list(st.session_state['checked_candidates'])).reset_index(drop=True)
                                st.session_state['df_candidates'] = df_candidates

                                # Clear the `checked_candidates` after ignoring them, so the checkboxes reset
                                st.session_state['checked_candidates'] = set()

                                st.success(f"Ignored {len(ignored_ids)} selected candidates.")
                                st.rerun()
                            else:
                                st.error("Failed to fetch job description.")
                else:
                    st.error("No candidates matched for the provided job description")
            else:
                st.error("Failed to fetch candidates.")


        mongo_client = MongoClient("mongodb://localhost:27017/")
        db = mongo_client['recruiter_ai']
        candidates_collection = db['candidates']

        if candidate_model.is_open():
            with candidate_model.container():
                selected_candidate_id = st.session_state.get('selected_candidate_id')
                candidate = candidates_collection.find_one({"ID": selected_candidate_id})

                if candidate:
                    candidate_name = candidate.get("Name")
                    st.markdown(f"<h2 style='text-align: center;'>Candidate: {candidate_name}</h2>", unsafe_allow_html=True)

                    # Define button colors based on candidate data
                    button_colors = {
                        "Message Sent/Call Queued": "green" if candidate.get('initiated_status') in ["Message Sent", "Call Queued", "Message Sent/Call Queued"] else "red",
                        "Interested": "green" if candidate.get('interested_status') == "Interested" else "red",
                        "Interview Scheduled": "green" if candidate.get('scheduled_status') == "Interview Scheduled" else "red",
                        "Interview Completed": "green" if candidate.get('completed_status') == "Interview Completed" else "red",
                        "Offer Sent": "green" if candidate.get('offer_status') == "Offer Sent" else "red"
                    }

                    # Scrollable container for buttons
                    st.markdown("""
                        <div style="display: flex; align-items: center; overflow-x: auto; overflow-y: hidden; max-height: 100px;">
                    """, unsafe_allow_html=True)

                    col1, col_arrow1, col2, col_arrow2, col3, col_arrow3, col4, col_arrow4, col5 = st.columns([1, 0.2, 1, 0.2, 1, 0.2, 1, 0.2, 1])

                    # Button rendering logic, update color immediately
                    def update_status(button_key, status_field, status_value):
                        candidates_collection.update_one(
                            {"ID": selected_candidate_id},
                            {"$set": {status_field: status_value}}
                        )
                        candidate[status_field] = status_value  # Update session state to reflect the status change

                    # Button 1: Message Sent/Call Queued
                    with col1:
                        st.markdown(f"""
                            <button style='background-color: {button_colors["Message Sent/Call Queued"]}; color: white; width: 100%; padding: 10px; border: none; border-radius: 5px;'>
                                Message Sent/Call Queued
                            </button>
                        """, unsafe_allow_html=True)

                    # Arrow 1
                    with col_arrow1:
                        st.markdown("<span style='font-size: 24px;'>→</span>", unsafe_allow_html=True)


                    # Button 2: Interested
                    is_message_sent_green = button_colors["Message Sent/Call Queued"] == "green"
                    with col2:
                        st.markdown(f"""
                            <button style='background-color: {button_colors["Interested"]}; color: white; width: 100%; padding: 10px; border: none; border-radius: 5px;'>
                                Interested
                            </button>
                        """, unsafe_allow_html=True)


                    # Arrow 2
                    with col_arrow2:
                        st.markdown("<span style='font-size: 24px;'>→</span>", unsafe_allow_html=True)


                    # Button 3: Interview Scheduled
                    with col3:
                        # Disable "Interview Scheduled" button if "Interested" is not green
                        is_interested_green = button_colors["Interested"] == "green"
                        is_button_disabled = not (is_message_sent_green and is_interested_green)

                        button_color = "green" if candidate.get("scheduled_status") == "Interview Scheduled" else "red"
                        button_key = f"interview_scheduled_{selected_candidate_id}"

                        if st.button("Interview Scheduled", key=button_key, disabled=is_button_disabled):
                            if not is_button_disabled:
                                update_status(button_key, "scheduled_status", "Interview Scheduled")

                        st.markdown(f"""
                            <button style='background-color: {button_color}; color: white; width: 100%; padding: 10px; border: none; border-radius: 5px;' { "disabled" if is_button_disabled else "" }>
                                Interview Scheduled
                            </button>
                        """, unsafe_allow_html=True)


                    # Arrow 3
                    with col_arrow3:
                        st.markdown("<span style='font-size: 24px;'>→</span>", unsafe_allow_html=True)

                    # Button 4: Interview Completed
                    with col4:

                        is_interview_scheduled_green = button_colors["Interview Scheduled"] == "green"
                        is_button_disabled = not (is_message_sent_green and is_interested_green and is_interview_scheduled_green)

                        button_color = "green" if candidate.get("completed_status") == "Interview Completed" else "red"
                        button_key = f"interview_completed_{selected_candidate_id}"

                        if st.button("Interview Completed", key=button_key, disabled=is_button_disabled):
                            if not is_button_disabled:
                                update_status(button_key, "completed_status", "Interview Completed")

                        st.markdown(f"""
                            <button style='background-color: {button_color}; color: white; width: 100%; padding: 10px; border: none; border-radius: 5px;' { "disabled" if is_button_disabled else "" }>
                                Interview Completed
                            </button>
                        """, unsafe_allow_html=True)


                    # Arrow 4
                    with col_arrow4:
                        st.markdown("<span style='font-size: 24px;'>→</span>", unsafe_allow_html=True)


                    # Button 5: Offer Sent
                    with col5:

                        # Check if "Interview Completed" is green
                        is_interview_completed_green = button_colors["Interview Completed"] == "green"
                        is_button_disabled = not (is_message_sent_green and is_interested_green and is_interview_scheduled_green and is_interview_completed_green)

                        button_color = "green" if candidate.get("offer_status") == "Offer Sent" else "red"
                        button_key = f"offer_sent_{selected_candidate_id}"

                        if st.button("Offer Sent", key=button_key, disabled=is_button_disabled):
                            if not is_button_disabled:
                                update_status(button_key, "offer_status", "Offer Sent")

                        st.markdown(f"""
                            <button style='background-color: {button_color}; color: white; width: 100%; padding: 10px; border: none; border-radius: 5px;' { "disabled" if is_button_disabled else "" }>
                                Offer Sent
                            </button>
                        """, unsafe_allow_html=True)


    def determine_interest(transcripts):
        """
        Determine if the candidate is interested based on the transcripts.
        Returns True if interested, False if disinterested, and None if unsure.
        """
        positive_keywords = ["interested", "yes", "looking for", "want", "job"]
        negative_keywords = ["not interested", "no interest", "decline", "disinterested", "no"]

        interest_found = None  # Start with None for undecided status

        for transcript in transcripts:
            text = transcript['text'].lower()
            
            # Check for positive keywords
            if any(keyword in text for keyword in positive_keywords):
                interest_found = True  # Update to True if any positive keyword is found

            # Check for negative keywords
            if any(keyword in text for keyword in negative_keywords):
                interest_found = False  # Update to False if any negative keyword is found

        return interest_found  # Return the final status


    def update_candidate_status(candidate_id, new_status, mongo_client):
        db = mongo_client['recruiter_ai']
        candidates_collection = db['candidates']

        candidate = candidates_collection.find_one({"ID": candidate_id})
        if candidate:
            # Determine the final status for initiated_status
            if "Message Sent" in candidate.get('initiated_status', '') and "Call Queued" in new_status:
                final_initiated_status = "Message Sent/Call Queued"
            elif "Message Sent" in new_status:
                final_initiated_status = "Message Sent"
            elif "Call Queued" in new_status:
                final_initiated_status = "Call Queued"
            else:
                final_initiated_status = new_status if new_status else "Not Sent"

            if "Interview Scheduled" in new_status:
                final_scheduled_status = "Interview Scheduled"
            elif "Call Queued" in candidate.get('scheduled_status', ''):
                final_scheduled_status = "Call Queued"
            else:
                final_scheduled_status = "Not Scheduled"

            if "Interview Completed" in new_status:
                final_completed_status = "Interview Completed"
            elif "Interview Scheduled" in candidate.get('scheduled_status', ''):
                final_completed_status = "Not Completed"
            else:
                final_completed_status = "Not Completed"

            if "Offer Sent" in new_status:
                final_offer_status = "Offer Sent"
            elif "Selected" in candidate.get('completed_status', ''):
                final_offer_status = "Not Offered"
            else:
                final_offer_status = "Not Offered"

            if "Interested" in new_status:
                final_interested_status = "Interested"
            elif "Disinterested" in new_status:
                final_interested_status = "Disinterested"
            else:
                final_interested_status = candidate.get('interested_status', 'Disinterested')

            # Update the candidate's status
            candidates_collection.update_one(
                {"ID": candidate_id},
                {
                    "$set": {
                        "initiated_status": final_initiated_status,
                        "scheduled_status": final_scheduled_status,
                        "completed_status": final_completed_status,
                        "offer_status": final_offer_status,
                        "interested_status": final_interested_status
                    },
                    "$unset": {
                        "Status": ""  # Remove the Status field
                    }
                }
            )

    def update_candidate_interest(candidate_id, mongo_client):
        db = mongo_client['recruiter_ai']
        candidates_collection = db['candidates']
        call_conversations_collection = db['call_conversations']

        candidate = candidates_collection.find_one({"ID": candidate_id})
        if not candidate:
            logger.info(f"Candidate with ID {candidate_id} not found.")
            return

        selected_job_id = candidate.get('selected_job_id')
        if not selected_job_id:
            logger.info(f"Selected Job ID for candidate ID {candidate_id} not found.")
            return

        # Fetch the most recent conversation for the selected job ID
        latest_conversation = call_conversations_collection.find_one(
            {"selected_job_id": selected_job_id},
            sort=[("timestamp", -1)]
        )

        if latest_conversation:
            interested = latest_conversation.get("interested", False)
            
            candidates_collection.update_one(
                {"ID": candidate_id},
                {"$set": {"interested_status": "Interested" if interested else "Not interested"}}
            )
            
            logger.info(f"Updated candidate ID {candidate_id} status to {'Interested' if interested else 'Not interested'}.")
        else:
            logger.info(f"No conversations found for candidate ID {candidate_id} and job ID {selected_job_id}.")
            
    
    def fetch_and_store_call_details(call_id, mongo_client):
        url = f"https://api.bland.ai/v1/calls/{call_id}"
        headers = {
            "Authorization": authorization_key,
            "Content-Type": "application/json"
        }

        while True:
            try:
                response = requests.get(url, headers=headers)
                response.raise_for_status()
                call_data = response.json()
                logger.info(f"API Response: {call_data}")

                if call_data.get("completed", False):
                    transcripts = call_data.get("transcripts", [])
                    interested = determine_interest(transcripts)

                    # Short response based on interest
                    if interested:
                        final_message = "Thank you, HR will contact you."
                    else:
                        final_message = "Thank you."
                        
                    call_conversation = {
                        "call_id": call_data.get("call_id"),
                        "timestamp": datetime.now(),
                        "caller": call_data.get("from", "Unknown"),
                        "receiver": call_data.get("to", "Unknown"),
                        "message": final_message,
                        "concatenated_transcript": call_data.get("concatenated_transcript"),
                        "call_length": call_data.get("call_length"),
                        "completed": call_data.get("completed"),
                        "started_at": call_data.get("started_at"),
                        "end_at": call_data.get("end_at"),
                        "transcripts": transcripts,
                        "interested": interested,
                        "selected_job_id": st.session_state.get('selected_job_id')
                    }

                    db = mongo_client['recruiter_ai']
                    collection = db.call_conversations

                    try:
                        insert_result = collection.insert_one(call_conversation)
                        logger.info(f"Call details stored successfully with _id: {insert_result.inserted_id}")

                        # Update candidate interest based on the latest call details
                        candidate_id = st.session_state.get('selected_candidate_id')
                        if candidate_id:
                            update_candidate_interest(candidate_id, mongo_client)

                    except Exception as e:
                        logger.error(f"Error inserting data into MongoDB: {str(e)}")
                    break
                else:
                    logger.info(f"Call {call_id} not completed yet. Retrying in 15 seconds...")
                    time.sleep(15)
            except requests.RequestException as e:
                logger.error(f"Request failed: {str(e)}")
                break

    def schedule_interviews():
        with st.spinner('Scheduling Interview for the selected Candidate...'):
            df_candidates = st.session_state.get('df_candidates')
            successful_updates = []
            error_messages = []

            user_id = st.session_state.get('user_id')
            user = fetch_user_data_by_id(user_id)
            api_token = user.get('whatsapp_api_token')
            cloud_number_id = user.get('whatsapp_cloud_number_id')

            if not api_token or not cloud_number_id:
                st.error("Missing API token or Cloud number ID.")
                st.session_state['setting_flag'] = True
                return

            mongo_client = pymongo.MongoClient("mongodb://localhost:27017/")

            for idx in st.session_state.get('checked_candidates', []):
                candidate = df_candidates.iloc[idx]
                phone_number = candidate['Phone']
                candidate_name = candidate['Name']
                candidate_email = candidate['Email']
                message_sent = False
                call_queued = False

                message_payload = {
                    "template_name": "hr_recruter_ai",
                    "language_code": "en_US",
                    "phone_number": phone_number,
                    "api_token": api_token,
                    "cloud_number_id": cloud_number_id,
                    "selected_job_id": st.session_state.get('selected_job_id')
                }

                try:
                    client_Email = EmailConfig_User()
                    client_Email.send_email_candidate(candidate_email, candidate_name, st.session_state.get('selected_job_id'))
                    message_response = requests.post("http://localhost:5000/send_message", json=message_payload)
                    message_response.raise_for_status()
                    if message_response.status_code == 200:
                        message_sent = True
                        logger.info(f"WhatsApp message sent successfully for candidate ID: {candidate['ID']}")
                    else:
                        error_messages.append(f"Failed to send WhatsApp message for candidate ID: {candidate['ID']} - Status code: {message_response.status_code}")
                        logger.error(f"Failed to send WhatsApp message for candidate ID: {candidate['ID']} - Status code: {message_response.status_code}")
                except requests.exceptions.ConnectionError as ce:
                    error_messages.append(f"Connection error occurred while sending WhatsApp message for candidate ID: {candidate['ID']} - {ce}")
                    logger.error(f"Connection error occurred while sending WhatsApp message for candidate ID: {candidate['ID']} - {ce}")
                except requests.exceptions.HTTPError as http_err:
                    error_messages.append(f"HTTP error occurred while sending WhatsApp message for candidate ID: {candidate['ID']} - {http_err}")
                    logger.error(f"HTTP error occurred while sending WhatsApp message for candidate ID: {candidate['ID']} - {http_err}")
                if not phone_number.startswith("+91"):
                    phone_number = "+91" + phone_number

                call_payload = {
                    "phone_number": phone_number,
                    "task": "Hello, I am from HR Recruiter. Are you interested in the job? If yes, HR will contact you",
                    # "max_duration": 0.20  # Set maximum call duration to 20 seconds
                }
                call_headers = {
                    "authorization": authorization_key,
                    "Content-Type": "application/json"
                }

                retry_attempts = 3
                for attempt in range(retry_attempts):
                    try:
                        call_response = requests.post("https://api.bland.ai/v1/calls", json=call_payload, headers=call_headers)
                        call_response.raise_for_status()
                        if call_response.status_code == 200:
                            call_response_json = call_response.json()
                            if call_response_json.get("status") == "success":
                                call_id = call_response_json.get("call_id")
                                if call_id:
                                    fetch_and_store_call_details(call_id, mongo_client)
                                    call_queued = True
                                    logger.info(f"Call API Response Status Code: 200")
                                    break
                                else:
                                    error_messages.append(f"Call scheduling failed for candidate ID: {candidate['ID']} - Call ID not returned.")
                                    logger.error(f"Call scheduling failed for candidate ID: {candidate['ID']} - Call ID not returned.")
                            else:
                                error_messages.append(f"Call scheduling failed for candidate ID: {candidate['ID']} - API response: {call_response_json}")
                                logger.error(f"Call scheduling failed for candidate ID: {candidate['ID']} - API response: {call_response_json}")
                        else:
                            error_messages.append(f"Failed to schedule call for candidate ID: {candidate['ID']} - Status code: {call_response.status_code}")
                            logger.error(f"Failed to schedule call for candidate ID: {candidate['ID']} - Status code: {call_response.status_code}")
                    except requests.exceptions.HTTPError as http_err:
                        error_messages.append(f"HTTP error occurred while scheduling call for candidate ID: {candidate['ID']} - {http_err}")
                        logger.error(f"HTTP error occurred while scheduling call for candidate ID: {candidate['ID']} - {http_err}")
                    time.sleep(15)  # Wait before retrying

                if not call_queued:
                    logger.info(f"Call not queued for candidate ID: {candidate['ID']}.")

                # Update status in MongoDB
                new_status = "Call Queued"
                if message_sent:
                    new_status = "Message Sent/Call Queued"
                else:
                    logger.info(f"Message not sent for candidate ID: {candidate['ID']}.")

                if call_queued:
                    df_candidates.at[idx, 'Status'] = new_status
                    candidate_id = int(candidate['ID'])
                    update_candidate_status(candidate_id, new_status, mongo_client)
                    update_candidate_interest(candidate_id, mongo_client)
                    successful_updates.append(idx)
                else:
                    error_messages.append(f"Call queuing failed for candidate ID: {candidate['ID']}.")
                    logger.error(f"Call queuing failed for candidate ID: {candidate['ID']}.")
        
            st.session_state['df_candidates'] = df_candidates
            st.session_state['checked_candidates'] = set()

            if successful_updates:
                st.success("Message sent and/or call queued successfully for the selected Candidate(s).")

            if error_messages:
                for message in error_messages:
                    st.error(message)

            if st.session_state.get('modal_open', False):
                candidate_id = st.session_state.get('modal_content')
                candidate_model(candidate_id)

            st.rerun()

    if st.session_state.get('whatsapp_success_flag'):
        st.success("Call queued successfully for the selected Candidate(s).")
        st.session_state['whatsapp_success_flag'] = False


    def fetch_user_data():
        url = 'http://localhost:8083/api/v1/user'
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()
        else:
            st.error("Failed to fetch user data.")
            return []
        
    def fetch_user_data_by_id(id):
        url = f'http://localhost:8083/api/v1/user/{id}'
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()
        else:
            st.error("Failed to fetch user data.")
            return []
        
    # Function to update user data
    def update_user_data(user_id, updated_user):
        url = f'http://localhost:8083/api/v1/user/{user_id}'
        response = requests.put(url, json=updated_user)
        if response.status_code == 200:
            #st.success("User updated successfully!")
            return True
        else:
            st.error("Failed to update user.")
            return False

    # Function to delete user
    def delete_user(user_id):
        url = f'http://localhost:8083/api/v1/user/{user_id}'
        response = requests.delete(url)
        if response.status_code == 200:
            #st.success("User deleted successfully!")
            collection.delete_many({"user_id": user_id})
            candidate_collection.delete_many({"user_id": user_id})
            return True
        else:
            st.error("Failed to delete user.")
            return False

    # Add User function
    def add_user(user):
        url = 'http://localhost:8083/api/v1/admin'
        response = requests.post(url, json=user)
        if response.status_code == 201:
            emailjson = response.json()
            email = emailjson.get("email")
            password = emailjson.get("password")
            client = EmailConfig_User()
            client.send_email(email,password)
            #st.success("User added successfully!")
            return True
        else:
            st.error("Failed to add user.")
            return False
    
    # Function to fetch member data
    def fetch_member_data():
        url = 'http://localhost:8083/api/v1/member'
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()
        elif response.status_code==404:
            st.warning("No members found")
            return False
        else:
            st.error("Failed to fetch member data.")
            return []
        
        
    # Function to update member data
    def update_member_data(member_id, updated_member):
        url = f'http://localhost:8083/api/v1/member/{member_id}'
        response = requests.put(url, json=updated_member)
        if response.status_code == 200:
            #st.success("User updated successfully!")
            return True
        else:
            st.error("Failed to update member.")
            return False

    # Function to delete member_id
    def delete_member(member_id):
        url = f'http://localhost:8083/api/v1/member/{member_id}'
        response = requests.delete(url)
        if response.status_code == 200:
            #st.success("User deleted successfully!")
            return True
        else:
            st.error("Failed to delete member.")
            return False

    # Add member function
    def add_member(member):
        url = 'http://localhost:8083/api/v1/member'
        response = requests.post(url, json=member)
        if response.status_code == 201:
            emailjson = response.json()
            email = emailjson.get("email")
            client = EmailConfig_User()
            client.send_email_Member(email)
            #st.success("Member added successfully and Email sent!")
            return True
        else:
            st.error("Failed to add Memeber.")
            return False

    def fetch_interview_rounds(selected_job_id):
        try:
            # Query the MongoDB collection for all documents with the given job ID
            interview_rounds = list(manage_job.find({"selected_job_id": selected_job_id}))

            # Check if any data is found
            if interview_rounds:
                return interview_rounds
            else:
                st.warning("No Interview Rounds Found for the selected job ID. Click on Add button to Add.")
                return []
        except Exception as e:
            st.error("Failed to fetch interview rounds")
            return []
        
    # Function to add an interview round
    def add_interview_round(current_job_description,selected_job_id, round_data):
        # Create a new job document with the initial interview round
        new_id = str(uuid.uuid4())
        round_data["current_job_description"] = current_job_description
        round_data["selected_job_id"] = selected_job_id
        round_data["id"] = new_id
        manage_job.insert_one(round_data)
        return True
    
    def update_job_data(selected_job_id, round_id, updated_round):
        # Find the job document with the specified selected_job_id and the specific interview round with round_id
        result = manage_job.update_one(
            {"selected_job_id": selected_job_id, "id": round_id},
            {"$set": {
                "name": updated_round.get("name", ""),
                "panel_members": updated_round.get("panel_members", 0),
                "panel_members_data": updated_round.get("panel_members_data", []),
                "interview_date": updated_round.get("interview_date", ""),
                "interview_time": updated_round.get("interview_time", "")
            }}
        )
        if result.modified_count > 0:
            return True
        else:
            return False
        
    def delete_job_panel(selected_job_id, panel_id):
        # First, check if the panel exists in the document
        result = manage_job.delete_one({"selected_job_id": selected_job_id, "id": panel_id})

        # Check if the deletion was successful
        if result.deleted_count > 0:
            return True
        else:
            return False

    # # Sidebar
    with st.sidebar:
        job_descriptions = fetch_job_descriptions()
        # Button to manage users if user is admin
        if st.session_state['user_role'] == 'admin':
            if st.button("Manage Users"):
                manage_users_modal.open()
            if st.button("Manage Members"):
                manage_members_modal.open()
            
        if st.button("Settings"):
            settings_model.open()
        if st.button("Logout", key="logout_button_home"):
            st.session_state.login_state = False
            st.session_state.login_successful = False
            st.session_state.message = ""
            st.session_state.login_inputs = {"email": "", "password": ""}

            st.session_state['current_job_description'] = ""
            st.session_state['selected_job_id'] = None
            st.session_state['modal_open'] = False
            st.session_state['modal_content'] = ""
            st.session_state['job_submitted'] = False
            st.session_state['job_updated'] = False
            st.session_state['view_candidates'] = False
            st.session_state['selected_candidates'] = []
            st.session_state['creating_new_job'] = False

            st.session_state.user_data = {}  # Clear user data
           
            
            st.rerun()

        # st.sidebar.button("Logout", on_click=logout)
        sac.divider(label='Job Description', align='center', color='gray')
        st.sidebar.button("New Job Description", on_click=new_job_description)
        # Search bar for job IDs
        search_term = st.sidebar.text_input("Search Job ID:", "")
        # Filter job descriptions based on search term
        filtered_jobs = [job for job in job_descriptions if search_term in str(job['_id'])]
        for job in filtered_jobs:
            job_id = job['_id']
            job_label = f"Job ID: {job_id}"
            is_selected = st.session_state['selected_job_id'] == job_id
            button_text = f"{job_label} {'(Selected)' if is_selected else ''}"

            if st.sidebar.button(button_text, key=f"job_{job_id}"):
                st.session_state['selected_job_id'] = job_id
                st.session_state['current_job_description'] = job.get('prompt', '')
                st.session_state['job_submitted'] = True
                st.session_state['job_updated'] = False
                st.session_state['creating_new_job'] = False
                st.session_state['edit_mode'] = False
                st.session_state['job_updated'] = False
                st.session_state['view_candidates'] = False
                st.session_state['checked_candidates'] = set()
                st.session_state['status_update'] = {}
                # Fetch the document from MongoDB
                job_document = collection.find_one({"_id": job_id})

                if job_document and job_document.get('first_save_clicked')  == True:
                    st.session_state['save_clicked'] = True
                else:
                    st.session_state['save_clicked'] = False
                    st.session_state['first_save_clicked'] = False

                st.rerun()

    # Job description input
    text_area_disabled = st.session_state['selected_job_id'] is not None
    job_description = st.text_area("Describe the Job Profile", value=st.session_state['current_job_description'], height=150, disabled=text_area_disabled)

    # Create columns for buttons
    col1, col2 = st.columns(2)

    with col1:
        submit_disabled = st.session_state['selected_job_id'] is not None
        if st.button("Submit", disabled=submit_disabled):
            submit_job_description(job_description)

    with col2:
        if st.session_state['job_submitted']:
            if st.session_state['selected_job_id']:
                col1, col2,col3 = st.columns(3)

                with col1:
                    open_modal = st.button("View Job Description")
                    with st.spinner('fetching job description...'):
                        if open_modal:
                            modal.open()

                with col2:
                    # Show the View Candidates button only after the first save
                    if st.session_state['first_save_clicked']:
                        view_candidates_button = st.button("View Candidates")
                        if view_candidates_button:
                            st.session_state['view_candidates'] = True
                    elif st.session_state['save_clicked']:
                        st.session_state['first_save_clicked'] = True
                        view_candidates_button = st.button("View Candidates")
                        if view_candidates_button:
                            st.session_state['view_candidates'] = True

                with col3:
                    if st.session_state['first_save_clicked']:
                        if st.button("Manage Job"):
                            manage_job_modal.open()
                    elif st.session_state['save_clicked']:
                        st.session_state['first_save_clicked'] = True
                        if st.button("Manage Job"):
                            manage_job_modal.open()

    if manage_job_modal.is_open():
        with manage_job_modal.container():
            # Show success message if the update was successful
            if 'update_success' in st.session_state and st.session_state.update_success:
                st.success("Panel updated successfully!")
                del st.session_state.update_success
            # Show success message if the update was successful
            if 'added_success' in st.session_state and st.session_state.added_success:
                st.success("Panel Added successfully!")
                del st.session_state.added_success
            # Show success message if the update was successful
            if 'delete_success' in st.session_state and st.session_state.delete_success:
                st.success("Panel Deleted successfully!")
                del st.session_state.delete_success

            if st.button("Add Round"):
                add_job_modal.open()

            selected_job_id = st.session_state['selected_job_id']

            progress_bar = st.progress(0)
            progress_bar.progress(10)  # Set initial progress
    
            # Fetch user data with loading indication
            interview_rounds = []
            try:
                interview_rounds = fetch_interview_rounds(selected_job_id)
                # Simulate loading time for demonstration purposes
                for i in range(20, 100, 10):
                    progress_bar.progress(i)
                    t.sleep(0.1)  # Simulate loading time
            except Exception as e:
                st.error("Error fetching Panel data")
    
            progress_bar.progress(100)
            progress_bar.empty()
    
            if interview_rounds:
                col1, col2, col3, col4, col5, col6, col7, col8 = st.columns([1, 2, 2, 2, 2, 2, 2, 2])
    
                with col1:
                    st.markdown("**id**")
                with col2:
                    st.markdown("**Name**")
                with col3:
                    st.markdown("**No of Panel Members**")
                with col4:
                    st.markdown("**Panel Members**")
                with col5:
                    st.markdown("**Date**")
                with col6:
                    st.markdown("**Time**")
                with col7:
                    st.markdown("**Edit**")
                with col8:
                    st.markdown("**Delete**")
    
            for index, panel in enumerate(interview_rounds):
                col1, col2, col3, col4, col5, col6, col7, col8 = st.columns([1, 2, 2, 2, 2, 2, 2, 2])
    
                with col1:
                    st.write(index + 1)
                with col2:
                    st.write(panel['name'])
                with col3:
                    st.write(panel['panel_members'])
                with col4:
                    st.markdown("\n".join(f"• {member}" for member in panel['panel_members_data']))
                with col5:
                    st.write(panel['interview_date'])
                with col6:
                    st.write(panel['interview_time'])
                with col7:
                    if st.button("Edit", key=f"edit_{index}"):
                        st.session_state['current_member_id'] = panel['id']
                        st.session_state['show_edit_dialog'] = True
                        confirm_edit_manage_job_dialog()
                with col8:
                    if st.button("Delete", key=f"delete_{index}"):
                        # Delete_Panel(panel['name'])
                        Delete_Panel(panel)

    # Define the content of the edit user modal
    if edit_job_modal.is_open():
        with edit_job_modal.container():
            # Load existing member data
            url = "http://localhost:8083/api/v1/member/"
            response = requests.get(url)
            if response.status_code == 200:
                member_data = response.json()
                member_names = [member["name"] for member in member_data]
            else:
                member_names = []

            member_id = st.session_state.get('current_member_id')

            # Get the current member to be edited
            rounds = fetch_interview_rounds(st.session_state['selected_job_id'])
    
            # Find the round with the given member ID
            member = next((round for round in rounds if round['id'] == member_id), None)
            if member:
                with st.form("edit_form"):
                    round_name = st.text_input("Round Name", value=member.get('name', ''))
                    panel_members = st.number_input("Number of Panel Members", min_value=1, step=1, value=member.get('panel_members', 1))
                    panel_members_data = st.multiselect("Add Members", options=member_names, default=member.get('panel_members_data', []))
                    # Handle interview date
                    today = datetime.now().date()
                    interview_date_str = member.get('interview_date', datetime.now().date().isoformat())
                    interview_date = datetime.fromisoformat(interview_date_str).date()
                    interview_date_input = st.date_input("Interview Date", value=interview_date,  min_value=today)

                    # Handle interview time
                    interview_time_str = member.get('interview_time', datetime.now().time().isoformat())
                    try:
                        interview_time = datetime.strptime(interview_time_str, "%H:%M:%S").time()
                    except ValueError:
                        interview_time = datetime.now().time()  # Default to current time if parsing fails
                    interview_time_input = st.time_input("Interview Time", value=interview_time)
                    
                    submit_button = st.form_submit_button("Update")

                    if submit_button:
                        if not round_name or not panel_members or not panel_members_data or not interview_date_input or not interview_time_input:
                            st.error("Please fill in all fields.")
                        elif round_name.strip() == "":
                            st.error("Round Name cannot be empty.")
                        elif len(panel_members_data) != panel_members:
                            st.error("The number of panel members does not match the selected members.")
                        else:
                            new_round = {
                                "name": round_name,
                                "panel_members": panel_members,
                                "panel_members_data": panel_members_data,
                                "interview_date": interview_date_input.isoformat(),
                                "interview_time": interview_time_input.isoformat()
                            }

                            # Compare the existing and updated user data to check for changes
                            changes_made = False
                            for key in new_round:
                                if new_round[key] != member.get(key):
                                    changes_made = True
                                    break
                                
                            if not changes_made:
                                st.info("No changes have been made.")
                            else:
                                update_result = update_job_data(st.session_state['selected_job_id'], member['id'], new_round)
                                logger.info(f"Update result: {update_result}")
                                if update_result:
                                    st.session_state.update_success = True
                                    edit_job_modal.close()
                                    st.rerun()
                                else:
                                    st.error("Failed to update the interview round. Please check the server logs for more details.")

    # panel modal
    if add_job_modal.is_open():
        with add_job_modal.container():
            
            url = "http://localhost:8083/api/v1/member/"
            response = requests.get(url)
            if response.status_code == 200:
                member_data = response.json()
                member_names = [member["name"] for member in member_data]
                member_email = [member["email"] for member in member_data]
                member_id = {member["name"]: member["id"] for member in member_data}
                members_info = {member["name"]: member["email"] for member in member_data}
                logger.info(f"Members info: {members_info}")
                logger.info(f"Member emails: {member_email}")
                logger.info(f"Member IDs: {member_id}")
            else:
                member_names = []

            with st.form("add_form"):
                round_name = st.text_input("Round Name")
                panel_members = st.number_input("Number of Panel Members",min_value=1, step=1, key="panel_members")
                panel_members_data = st.multiselect("Add Members",options=member_names, key="panel_members_data")
                today = datetime.now().date()
                interview_date = st.date_input("Interview Date", min_value=today, key="interview_date")
                interview_time = st.time_input("Interview Time", key="interview_time")

                if st.form_submit_button("Add Panel"):
                    with st.spinner('Loading...'):
                        # Validate inputs
                        if not round_name or not panel_members or not panel_members_data or not interview_date or not interview_time:
                            st.error("Please fill in all fields.")
                        elif round_name.strip() == "":
                            st.error("Round Name cannot be empty.")
                        elif len(panel_members_data) != panel_members:
                            st.error("The number of panel members does not match the selected members.")
                        else:
                            new_round = {
                                "name": round_name,
                                "panel_members": panel_members,
                                "panel_members_data": panel_members_data,
                                "panel_member_name" : panel_members_data,
                                "member_id": member_id,
                                "interview_date": interview_date.isoformat(),
                                "interview_time": interview_time.isoformat(),
                            }
                            # Add the interview round
                            if add_interview_round(st.session_state['current_job_description'],st.session_state['selected_job_id'], new_round):
                                emails = []
                                for name in panel_members_data:
                                    if name in members_info:
                                        email = members_info[name]
                                        emails.append(email)
                                        job_id = st.session_state['selected_job_id']
                                        logger.info(f"Email found: {email}")
                                        client = EmailConfig_User()
                                        client.send_email_SchdeulePanel(email, job_id, interview_date, interview_time, panel_members_data) 
                                    else:
                                        logger.warning(f"No email found for {name}")
                                st.session_state.added_success = True
                                add_job_modal.close()
                                st.rerun()
                            else:
                                st.error("Failed to add the interview round.")
                    
 
    if modal.is_open():
        with modal.container():
            if st.session_state['update_success_flag']:
                st.success("Job description updated successfully!")
                st.session_state['update_success_flag'] = False

            if st.session_state['update_error_flag']:
                st.error("Failed to update job description. Please try again.")
                st.session_state['update_error_flag'] = False

            if st.session_state['update_fetch_error_flag']:
                st.error("Failed to fetch job description. Please try again.")
                st.session_state['update_fetch_error_flag'] = False

            if st.session_state['update_warning_flag']:
                st.warning("No job description selected.")
                st.session_state['update_warning_flag'] = False

            if st.session_state['selected_job_id'] is not None:
                api_url = f"http://localhost:8081/api/v1/jd/{st.session_state['selected_job_id']}"
                response = requests.get(api_url)

                if response.status_code == 200:
                    jd_response = response.json()
                    # Determine if text area should be disabled
                    is_disabled = not st.session_state['edit_mode'] and st.session_state['job_submitted']
                    job_description = st.text_area("Job Description", value=jd_response.get('job_description', ''), height=250, disabled=is_disabled)
                    col1, col2 = st.columns(2)
                    with col1:
                        # Add Save button
                        save_button = st.button("Resource Candidates", disabled=st.session_state['save_clicked'])
                        with st.spinner('Filtering Candidates...'):
                            if save_button:
                                # Save job description logic
                                st.session_state['first_save_clicked'] = True
                                candidate_url = "http://localhost:8085/api/v1/resource/"
                                payload = {"job_description": job_description}  # JD from Streamlit database
                                
                                try:
                                    candidate_response = requests.post(candidate_url, json=payload)
                                    if candidate_response.status_code == 200:
                                        candidate_data = candidate_response.json()
                                        resource_id = candidate_data.get("id")
                                        selected_job_id = st.session_state['selected_job_id']
                                        user_id = st.session_state.get('user_id')

                                        # Check if the resource_id already exists in the database
                                        existing_record = collection.find_one({"_id": selected_job_id, "resource_id": {"$exists": True}})

                                        if existing_record:
                                            # Update the existing record with the new resource_id
                                            collection.update_one(
                                                {"_id": selected_job_id},
                                                {"$set": {"resource_id": resource_id}}
                                            )
                                        else:
                                            # Insert resource_id into the corresponding JD in the database
                                            collection.update_one(
                                                {"_id": selected_job_id},
                                                {"$set": {"resource_id": resource_id}}
                                            )

                                        candidate_list = []
                                        if "resource" in candidate_data:
                                            resource_data = candidate_data["resource"]
                                            if resource_data != []:
                                                if resource_data:
                                                    for item in resource_data:
                                                        try:
                                                            # Parse the item string to a list of candidate information
                                                            outer_list = ast.literal_eval(item)
                                                            for candidate_info in outer_list:
                                                                candidate_dict = {
                                                                    "ID": candidate_info[0],
                                                                    "Name": candidate_info[1],
                                                                    "Email": candidate_info[2],
                                                                    "Phone": candidate_info[3],
                                                                    "Status": "Interview Not Scheduled",
                                                                    "resource_id": resource_id,
                                                                    "selected_job_id": selected_job_id,
                                                                    "user_id": user_id
                                                                }
                                                                candidate_list.append(candidate_dict)
                                                            st.success("Candidate filtered Successfully...")
                                                            selected_job_id = st.session_state['selected_job_id'] 
                                                            first_save_clicked = True
                                                            collection.update_one({"_id": selected_job_id},{"$set": {"first_save_clicked": first_save_clicked}})
                                                            # st.session_state['save_clicked']=True
                                                        except Exception as e:
                                                            st.error("No candidates matched for the Job Description...")
                                            else:
                                                st.error("No candidates matched for the Job Description...")

                                        # Save or update candidates in MongoDB
                                        if candidate_list:
                                            # Define the query to find candidates with the same job ID and user ID
                                            query = {"selected_job_id": selected_job_id, "user_id": user_id}
                                            # Delete existing candidates with the same job ID and user ID
                                            candidate_collection.delete_many(query)
                                            candidate_collection.insert_many(candidate_list)

                                    else:
                                        st.error(f"Error: {candidate_response.status_code}")
                                        st.write(candidate_response.json())
                                except requests.exceptions.RequestException as e:
                                    st.error("Request failed")                                

                    with col2:
                    # Create columns for Edit and Update buttons
                        download_col,edit_col, update_col = st.columns([1,1, 1])
                        with download_col:
                            download_button = st.button("Download")
                            if download_button:
                                handle_download()

                        with edit_col:
                            if st.button("Edit"):
                                confirm_edit_dialog()
                           
                        with update_col:
                            if st.session_state['edit_mode']:
                                update_button = st.button("Update")
                                if update_button:
                                    update_job_description(job_description)
                                    st.session_state['edit_mode'] = False 
                                    st.session_state['job_updated'] = True
                                    st.session_state['save_clicked']=False  
                                    st.rerun()
                            
                else:
                    st.session_state['update_fetch_error_flag'] = True
            else:
                st.session_state['update_warning_flag'] = True

    if st.session_state['view_candidates']:
        display_candidates()

    # Define the content of the manage users modal
    if manage_users_modal.is_open():
        with manage_users_modal.container():
            # Show success message if the update was successful
            if 'update_success' in st.session_state and st.session_state.update_success:
                st.success("User updated successfully!")
                del st.session_state.update_success
            # Show success message if the update was successful
            if 'added_success' in st.session_state and st.session_state.added_success:
                st.success("User Added successfully and Email sent!")
                del st.session_state.added_success
            # Show success message if the update was successful
            if 'delete_success' in st.session_state and st.session_state.delete_success:
                st.success("User Deleted successfully!")
                del st.session_state.delete_success
            # Add User Button
            if st.button("Add User"):
                add_user_modal.open()

            # Display progress bar
            progress_bar = st.progress(0)
            progress_bar.progress(10)  # Set initial progress
    
            # Fetch user data with loading indication
            user_data = []
            try:
                user_data = fetch_user_data()
                # Simulate loading time for demonstration purposes
                for i in range(20, 100, 10):
                    progress_bar.progress(i)
                    t.sleep(0.1)  # Simulate loading time
            except Exception as e:
                st.error("Error fetching user data")
    
            progress_bar.progress(100)
            progress_bar.empty()
            # Complete progress bar
            
            if user_data:
                # Filter out the admin details
                filtered_user_data = [user for user in user_data if user['email'] != 'admin@gmail.com']

                col1, col2, col3, col4,col5 = st.columns([3, 2, 3, 2, 2])

                with col1:
                    st.markdown("**id**")
                with col2:
                    st.markdown("**Name**")
                with col3:
                    st.markdown("**email**")
                with col4:
                    st.markdown("**Edit**")
                with col5:
                    st.markdown("**Delete**")
    
                for index, user in enumerate(filtered_user_data):
                    col1, col2, col3, col4, col5 = st.columns([3, 2, 3, 2, 2])

                    with col1:
                        st.write(index + 1)
                    with col2:
                        st.write(user['name'])
                    with col3:
                        st.write(user['email'])
                    with col4:
                        if st.button("Edit", key=f"edit_{index}"):
                            st.session_state['current_user'] = user
                            # edit_user_modal.open()
                            st.session_state['show_edit_dialog'] = True
                            confirm_edit_user_dialog()  # Open the confirmation dialog
                    with col5:
                        if st.button("Delete", key=f"delete_{index}"):
                            # Delete_Users(user['name'])
                            Delete_Users(user)
        
    # Define the content of the edit user modal
    if edit_user_modal.is_open():
        with edit_user_modal.container():
            user = st.session_state.get('current_user', {})

            if user:
                with st.form("edit_form"):
                    name = st.text_input("Name", value=user.get('name', ''))
                    email = st.text_input("Email", value=user.get('email', ''))
                    mobile_number = st.text_input("Mobile Number", value=user.get('mobile_number', ''))
                    location = st.text_input("Location", value=user.get('location', ''))
                    

                    submit_button = st.form_submit_button("Update")
                    if submit_button:
                        # Validate inputs
                        if not name or not email or not mobile_number or not location:
                            st.error("Please fill in all fields.")
                        elif len(name) > 20:
                            st.error("Name should be up to 20 characters.")
                        elif not validate_name(name):
                            st.error("Name should start with Uppercase")
                        elif not validate_email(email):
                            st.error("Invalid email format.")
                        elif not validate_mobile_number(mobile_number):
                            st.error("Mobile number must be exactly 10 digits.")
                        else:
                            updated_user = {
                                "name": name,
                                "email": email,
                                "mobile_number": int(mobile_number),
                                "location": location
                            }

                            # Compare the existing and updated user data to check for changes
                            changes_made = False
                            for key in updated_user:
                                if updated_user[key] != user.get(key):
                                    changes_made = True
                                    break

                            if not changes_made:
                                st.info("No changes have been made.")
                            else:
                                if update_user_data(user['id'], updated_user):
                                    st.session_state.update_success = True
                                    edit_user_modal.close()
                                    st.rerun()

    # Define the content of the add user modal
    if add_user_modal.is_open():
        with add_user_modal.container():

            with st.form("add_form"):
                name = st.text_input("Name")
                email = st.text_input("Email")
                mobile_number = st.text_input("Mobile Number")
                location = st.text_input("Location")

                if st.form_submit_button("Add User"):
                    with st.spinner('Loading...'):
                        # Validate inputs
                        if not name or not email or not mobile_number or not location:
                            st.error("Please fill in all fields.")
                        elif len(name) > 20:
                            st.error("Name should be up to 20 characters.")
                        elif not validate_name(name):
                            st.error("Name should start with Uppercase")
                        elif not validate_email(email):
                            st.error("Invalid email format.")
                        elif not validate_mobile_number(mobile_number):
                            st.error("Mobile number must be exactly 10 digits.")
                        else:
                            new_user = {
                                "name": name,
                                "email": email,
                                "mobile_number": int(mobile_number),
                                "location": location
                            }
                        if add_user(new_user):
                            st.session_state.added_success = True
                            add_user_modal.close()
                            st.rerun()
                        else:
                            st.session_state.added_success=False
                        
    # Define the content of the manage members modal
    if manage_members_modal.is_open():
        with manage_members_modal.container():
            # Show success message if the update was successful
            if 'update_success' in st.session_state and st.session_state.update_success:
                st.success("Member updated successfully!")
                del st.session_state.update_success
            # Show success message if the update was successful
            if 'added_success' in st.session_state and st.session_state.added_success:
                st.success("Member Added successfully and Email sent!")
                del st.session_state.added_success
            # Show success message if the update was successful
            if 'delete_success' in st.session_state and st.session_state.delete_success:
                st.success("Member Deleted successfully!")
                del st.session_state.delete_success
            # Add Member Button
            if st.button("Add Member"):
                add_members_modal.open()

            # Display progress bar
            progress_bar = st.progress(0)
            progress_bar.progress(10)  # Set initial progress
    
            # Fetch member data with loading indication
            member_data = []
            try:
                member_data = fetch_member_data()
                # Simulate loading time for demonstration purposes
                for i in range(20, 100, 10):
                    progress_bar.progress(i)
                    t.sleep(0.1)  # Simulate loading time
            except Exception as e:
                st.error("Error fetching member data")
    
            progress_bar.progress(100)
            progress_bar.empty()
            # Complete progress bar
            
            if member_data:
                col1, col2, col3, col4,col5 = st.columns([3, 2, 3, 2, 2])

                with col1:
                    st.markdown("**id**")
                with col2:
                    st.markdown("**Name**")
                with col3:
                    st.markdown("**email**")
                with col4:
                    st.markdown("**Edit**")
                with col5:
                    st.markdown("**Delete**")
    
                for index, member in enumerate(member_data):
                    col1, col2, col3, col4, col5 = st.columns([3, 2, 3, 2, 2])

                    with col1:
                        st.write(index + 1)
                    with col2:
                        st.write(member['name'])
                    with col3:
                        st.write(member['email'])
                    with col4:
                        if st.button("Edit", key=f"edit_{index}"):
                            st.session_state['current_member'] = member
                            st.session_state['show_edit_dialog'] = True
                            confirm_edit_member_dialog()
                    with col5:
                        if st.button("Delete", key=f"delete_{index}"):
                            Delete_Members(member)
        
    # Define the content of the edit member modal
    if edit_members_modal.is_open():
        with edit_members_modal.container():
            member = st.session_state.get('current_member', {})

            if member:
                with st.form("edit_form"):
                    name = st.text_input("Name", value=member.get('name', ''))
                    email = st.text_input("Email", value=member.get('email', ''))
                    mobile_number = st.text_input("Mobile Number", value=member.get('mobile_number', ''))
                    location = st.text_input("Location", value=member.get('location', ''))
                    

                    submit_button = st.form_submit_button("Update")
                    if submit_button:
                        # Validate inputs
                        if not name or not email or not mobile_number or not location:
                            st.error("Please fill in all fields.")
                        elif len(name) > 20:
                            st.error("Name should be up to 20 characters.")
                        elif not validate_name(name):
                            st.error("Name should start with Uppercase")
                        elif not validate_email(email):
                            st.error("Invalid email format.")
                        elif not validate_mobile_number(mobile_number):
                            st.error("Mobile number must be exactly 10 digits.")
                        else:
                            updated_member = {
                                "name": name,
                                "email": email,
                                "mobile_number": int(mobile_number),
                                "location": location
                            }

                            # Compare the existing and updated member data to check for changes
                            changes_made = False
                            for key in updated_member:
                                if updated_member[key] != member.get(key):
                                    changes_made = True
                                    break

                            if not changes_made:
                                st.info("No changes have been made.")
                            else:
                                if update_member_data(member['id'], updated_member):
                                    st.session_state.update_success = True
                                    edit_members_modal.close()
                                    st.rerun()

    # Define the content of the add member modal
    if add_members_modal.is_open():
        with add_members_modal.container():

            with st.form("add_form"):
                name = st.text_input("Name")
                email = st.text_input("Email")
                mobile_number = st.text_input("Mobile Number")
                location = st.text_input("Location")

                if st.form_submit_button("Add Member"):
                    with st.spinner('Loading...'):
                        # Validate inputs
                        if not name or not email or not mobile_number or not location:
                            st.error("Please fill in all fields.")
                        elif len(name) > 20:
                            st.error("Name should be up to 20 characters.")
                        elif not validate_name(name):
                            st.error("Name should start with Uppercase")
                        elif not validate_email(email):
                            st.error("Invalid email format.")
                        elif not validate_mobile_number(mobile_number):
                            st.error("Mobile number must be exactly 10 digits.")
                        else:
                            new_member = {
                                "name": name,
                                "email": email,
                                "mobile_number": int(mobile_number),
                                "location": location
                            }
                        if add_member(new_member):
                            st.session_state.added_success = True
                            add_members_modal.close()
                            st.rerun()
                        else:
                            st.session_state.added_success=False
                        
    # Define the content of the add user modal
    if settings_model.is_open():
        with settings_model.container():
            # Show success message if the update was successful
            if 'update_success' in st.session_state and st.session_state.update_success:
                st.success("User updated successfully!")
                del st.session_state.update_success
            # Fetch user data
            user_id = st.session_state.get('user_id')
            user = fetch_user_data_by_id(user_id)  # Fetch user data by ID
    
            if user:
                with st.form("edit_form"):
                    # Form fields pre-populated with user data
                    name = st.text_input("Name", value=user.get('name', ''))
                    email = st.text_input("Email", value=user.get('email', ''), disabled=True)
                    mobile_number = st.text_input("Mobile Number", value=str(user.get('mobile_number', '')))  # Convert to string for text input
                    location = st.text_input("Location", value=user.get('location', ''))
                    # Password fields for resetting password
                    new_password = st.text_input("New Password", value="", type="password")
                    confirm_password = st.text_input("Confirm New Password", value="", type="password")
                    whatsapp_api_token = st.text_input("WhatsApp API Token", value=user.get('whatsapp_api_token', ''))
                    whatsapp_cloud_number_id = st.text_input("WhatsApp Cloud Number", value=user.get('whatsapp_cloud_number_id', ''))
                    
                    submit_button = st.form_submit_button("Update")
                    if submit_button:
                        try:
                            # Validate inputs
                            if not name or not email or not mobile_number or not location:
                                st.error("Please fill in all fields.")
                            elif len(name) > 20:
                                st.error("Name should be up to 20 characters.")
                            elif not validate_name(name):
                                st.error("Name should start with Uppercase")
                            elif not validate_email(email):
                                st.error("Invalid email format.")
                            elif new_password and not validate_password(new_password):
                                st.error("Password must be at least 8 characters long and include uppercase, lowercase,special character and a number.")
                            elif new_password != confirm_password:
                                st.error("New password and confirmation do not match.")
                            elif not validate_mobile_number(mobile_number):
                                st.error("Mobile number must be exactly 10 digits.")
                            else:
                                # Convert mobile number to int, handle errors if conversion fails
                                updated_user = {
                                    "name": name,
                                    "email": email,
                                    "mobile_number": int(mobile_number),
                                    "location": location,
                                    "whatsapp_api_token": whatsapp_api_token,
                                    "whatsapp_cloud_number_id": whatsapp_cloud_number_id
                                }
                                # Include the new password if provided
                                if new_password:
                                    updated_user["password"] = new_password
                                # Compare the existing and updated user data to check for changes
                                changes_made = False
                                for key in updated_user:
                                    if key == 'password':
                                        if new_password and updated_user[key] != user.get(key):
                                            changes_made = True
                                            break
                                    elif updated_user[key] != user.get(key):
                                        changes_made = True
                                        break
                                    
                                if not changes_made:
                                    st.info("No changes have been made.")
                                else:
                                    if update_user_data(user_id, updated_user):
                                        st.session_state.update_success = True
                                        st.rerun()  # Reload the page to reflect changes
                        except ValueError:
                            st.error("Please enter a valid mobile number.")
                            
if __name__ == "__main__":
    main()